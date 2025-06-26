import os
import re
import zipfile
from docx import Document
import xml.etree.ElementTree as ET

def extract_images_with_precise_index(docx_path, output_folder="images"):
    """Trích xuất ảnh và gắn tên theo chỉ mục gần nhất như Bài 1.23, Hình 1.1 hoặc tiêu đề chương"""

    # Tạo thư mục lưu ảnh
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Đọc document bằng python-docx
    doc = Document(docx_path)

    # ✅ SỬA: LẤY TÊN FILE GỐC VÀ CẮT TRƯỚC CHỮ "RUOT"
    base_filename = os.path.splitext(os.path.basename(docx_path))[0]
    
    # Tìm vị trí của "ruot" hoặc "Ruot" (không phân biệt hoa thường)
    ruot_match = re.search(r'\s*[rR]uot', base_filename)
    if ruot_match:
        base_filename = base_filename[:ruot_match.start()].strip()
    
    base_filename = base_filename + "- KNTT "
    print(f"📁 Tên file gốc (đã cắt): {base_filename}")

    # Bước 1: Mapping relationship ID -> file ảnh
    image_files = {}
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        try:
            rels_content = docx_zip.read('word/_rels/document.xml.rels').decode('utf-8')
            root = ET.fromstring(rels_content)

            for rel in root.findall('.//{*}Relationship'):
                if rel.get('Target', '').startswith('media/'):
                    image_files[rel.get('Id')] = rel.get('Target')
        except Exception as e:
            print(f"❌ Không thể đọc relationships: {e}")
            return
    
    print(f"📋 Tìm thấy {len(image_files)} file ảnh trong document")

    # Bước 2: Duyệt đoạn văn + bảng, tìm ảnh và chỉ mục gần nhất
    images_to_save = []
    current_index = None
    current_title = None  # ✅ THÊM: Lưu tiêu đề hiện tại

    def is_title_paragraph(para):
        """Kiểm tra xem paragraph có phải là tiêu đề không"""
        text = para.text.strip()
        
        # ✅ PATTERN NHẬN DIỆN TIÊU đề
        title_patterns = [
            r'^[A-Z][^a-z]*$',  # Toàn chữ hoa: "PHÉP NHÂN VÀ PHÉP CHIA SỐ TỰ NHIÊN"
            r'^[A-Z]\.\s*[A-Z]',  # "A. KIẾN THỨC CẦN NHỚ"
            r'^[IVX]+\.\s*[A-Z]',  # "I. GIỚI THIỆU"
            r'^\d+\.\s*[A-Z]',     # "1. CHƯƠNG ĐẦU"
            r'^BÀI\s*$',           # Chỉ có chữ "BÀI"
            r'^CHƯƠNG\s+[IVX\d]',  # "CHƯƠNG I", "CHƯƠNG 1"
        ]
        
        # Kiểm tra độ dài (tiêu đề thường ngắn hơn 100 ký tự)
        if len(text) > 100:
            return False
            
        # Kiểm tra pattern
        for pattern in title_patterns:
            if re.match(pattern, text):
                return True
                
        # Kiểm tra style của paragraph (nếu có)
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if any(keyword in style_name for keyword in ['heading', 'title', 'header']):
                return True
                
        return False

    def clean_title(title):
        """Làm sạch tiêu đề để dùng làm tên file"""
        # Loại bỏ ký tự đặc biệt
        cleaned = re.sub(r'[<>:"/\\|?*]', '', title)
        # Thay thế khoảng trắng liền nhau bằng 1 khoảng trắng
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        # Giới hạn độ dài
        if len(cleaned) > 50:
            cleaned = cleaned[:47] + "..."
        return cleaned

    def process_paragraphs(paragraphs, context_text=""):
        nonlocal current_index, current_title
        paragraph_count = 0
        last_index_paragraph = -1

        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            paragraph_count += 1
            
            # ✅ KIỂM TRA TIÊU ĐỀ - ĐỒNG BỘ VỚI PREVIEW
            if is_title_paragraph(para):
                current_title = clean_title(text)
                current_index = None  # Reset như preview
                print(f"📋 Tiêu đề mới: {text}")
                print(f"    → Reset current_index, lưu title: '{current_title}'")
                continue
            
            # ✅ SỬA: DÙNG LOGIC GIỐNG HỆT PREVIEW
            found_index = None
            
            # 1. TÌM HÌNH TRƯỚC (giống preview)
            hinh_matches = re.findall(r'Hình\s+(\d+\.\d+)', text, re.IGNORECASE)
            if hinh_matches:
                found_index = f"Hình {hinh_matches[-1]}"  # Lấy số cuối cùng
                print(f"🖼️  Phát hiện hình: {found_index}")

            if not found_index:
                h_matches = re.findall(r'H\.(\d+\.\d+)', text)
                if h_matches:
                    found_index = f"Hình {h_matches[-1]}"  # Chuyển H.4.22 → Hình 4.22
                    print(f"🖼️  Phát hiện hình (rút gọn): H.{h_matches[-1]} → {found_index}")
            # 2. NẾU KHÔNG CÓ HÌNH, TÌM SỐ BÀI (giống preview)
            if not found_index:
                bai_matches = re.findall(r'(\d+\.\d+)', text)
                if bai_matches:
                    found_index = f"Bài {bai_matches[-1]}"  # Lấy số cuối cùng
                    print(f"📚 Phát hiện bài: {found_index}")
        
            if found_index:
                current_index = found_index
                last_index_paragraph = paragraph_count
                continue

            # ✅ XỬ LÝ ẢNH - LOGIC GIỐNG PREVIEW
            for run in para.runs:
                blips = run._element.xpath('.//a:blip')
                for blip in blips:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id in image_files:
                        
                        # ✅ LOGIC GIỐNG HỆT PREVIEW: current_index → current_title → "Không xác định"
                        final_index = current_index or current_title or "Không xác định"
                        
                        # Xác định lý do (cho debug)
                        if current_index:
                            reason = "current_index"
                        elif current_title:
                            reason = "current_title"
                        else:
                            reason = "không xác định"
                        
                        print(f"    🖼️  Ảnh → gắn với '{final_index}' (lý do: {reason})")
                        
                        images_to_save.append({
                            'file_path': image_files[embed_id],
                            'index': final_index,
                            'context': (text or context_text)[:50] + "..."
                        })

    # Duyệt đoạn văn chính
    print("\n🔍 Đang duyệt paragraphs...")
    process_paragraphs(doc.paragraphs)

    # Duyệt bảng
    print("🔍 Đang duyệt tables...")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    process_paragraphs(cell.paragraphs, context_text=cell.text.strip())

    # Bước 3: Lưu ảnh theo format mới
    print(f"\n🖼️ Tìm thấy {len(images_to_save)} ảnh:")

    if len(images_to_save) == 0:
        print("❌ Không tìm thấy ảnh nào để lưu.")
        return

    saved_count = 0

    # ✅ ĐẾM STT THEO THỨ TỰ XUẤT HIỆN
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        for i, img in enumerate(images_to_save, 1):
            try:
                # Đọc dữ liệu ảnh
                image_data = docx_zip.read(f"word/{img['file_path']}")
                
                # Xác định extension từ file gốc
                _, ext = os.path.splitext(img['file_path'])
                ext = ext.lower() or ".png"

                # ✅ TẠO TÊN FILE MỚI: STT + tên file gốc + bài/hình/title
                index = img['index'] or "Không xác định"
                
                filename = f"{i:02d} - {base_filename} - {index}{ext}"
                
                # Lưu file
                filepath = os.path.join(output_folder, filename)
                with open(filepath, 'wb') as f:
                    f.write(image_data)

                print(f"✅ {filename}")
                if img['context'].strip():
                    print(f"    Context: {img['context']}")
                
                saved_count += 1

            except Exception as e:
                print(f"❌ Lỗi khi lưu ảnh {i}: {e}")

    # Báo cáo kết quả
    print(f"\n🎉 Đã lưu {saved_count}/{len(images_to_save)} ảnh vào '{output_folder}'")
    
    if saved_count > 0:
        indices = [img['index'] for img in images_to_save if img['index']]
        unique_indices = sorted(set(indices), key=lambda x: (x.startswith('Bài'), x.startswith('Hình'), x))
        print(f"📊 Các chỉ mục được sử dụng: {unique_indices}")

def preview_images_and_indices(docx_path):
    """Xem trước danh sách ảnh và chỉ mục mà không lưu ảnh"""
    
    doc = Document(docx_path)
    base_filename = os.path.splitext(os.path.basename(docx_path))[0]
    current_index = None
    current_title = None
    image_count = 0
    
    print("=" * 70)
    print("🔍 XEM TRƯỚC ẢNH VÀ CHỈ MỤC")
    print("=" * 70)
    print(f"📁 Tên file gốc: {base_filename}")
    print()

    def is_title_paragraph(para):
        """Hàm kiểm tra tiêu đề tương tự như trên"""
        text = para.text.strip()
        title_patterns = [
            r'^[A-Z][^a-z]*$',
            r'^[A-Z]\.\s*[A-Z]',
            r'^[IVX]+\.\s*[A-Z]',
            r'^\d+\.\s*[A-Z]',
            r'^BÀI\s*$',
            r'^CHƯƠNG\s+[IVX\d]',
        ]
        
        if len(text) > 100:
            return False
            
        for pattern in title_patterns:
            if re.match(pattern, text):
                return True
                
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if any(keyword in style_name for keyword in ['heading', 'title', 'header']):
                return True
                
        return False

    def clean_title(title):
        """Hàm làm sạch tiêu đề tương tự như trên"""
        cleaned = re.sub(r'[<>:"/\\|?*]', '', title)
        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
        if len(cleaned) > 50:
            cleaned = cleaned[:47] + "..."
        return cleaned

    def process(paragraphs):
        nonlocal current_index, current_title, image_count

        for para in paragraphs:
            text = para.text.strip()
            
            # Kiểm tra tiêu đề
            if is_title_paragraph(para):
                current_title = clean_title(text)
                print(f"📋 Title: {current_title}")
                continue
            
            # Tìm chỉ mục theo thứ tự ưu tiên
            found_index = None
            
            hinh_matches = re.findall(r'Hình\s+(\d+\.\d+)', text, re.IGNORECASE)
            if hinh_matches:
                found_index = f"Hình {hinh_matches[-1]}"
            
            if not found_index:
                bai_matches = re.findall(r'(\d+\.\d+)', text)
                if bai_matches:
                    found_index = f"Bài {bai_matches[-1]}"
            
            if found_index:
                current_index = found_index

            # Kiểm tra ảnh
            for run in para.runs:
                if run._element.xpath('.//a:blip'):
                    image_count += 1
                    
                    # ✅ LOGIC MỚI: Ưu tiên current_index → current_title → "Không xác định"
                    index_display = current_index or current_title or "Không xác định"
                    filename_preview = f"{image_count:02d} - {base_filename} - {index_display}.png"
                    
                    display_text = text[:50] + "..." if len(text) > 50 else text
                    if not display_text.strip():
                        display_text = "[Paragraph chỉ có ảnh]"
                    
                    print(f"🖼️  Ảnh {image_count:2d}: {filename_preview}")
                    print(f"    📝 Context: {display_text}")
                    print()

    # Xử lý document
    process(doc.paragraphs)

    # Xử lý tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    process(cell.paragraphs)

    print("=" * 70)
    print(f"📊 Tổng cộng: {image_count} ảnh sẽ được trích xuất")
    print("=" * 70)

def main():
    """Hàm main với menu"""
    
    print("🎯 CHƯƠNG TRÌNH TRÍCH XUẤT ẢNH TỪ DOCX")
    print("🏷️  Format: STT - Tên_file - Bài/Hình/Tiêu_đề")
    print("🧠 Ưu tiên: Số bài → Hình → Tiêu đề chương")
    print("=" * 60)
    
    # Đường dẫn mặc định
    default_path = r"E:\Data\Work\SeperateImage\output\SBT Toan 6 tap 1 ruot(TB2025)_KNTT (14.3.2025) (1)_converted.docx"
    
    docx_path = input("Nhập đường dẫn file DOCX (Enter cho mặc định): ").strip()
    if not docx_path:
        docx_path = default_path
        print(f"📁 Sử dụng file mặc định")

    if not os.path.exists(docx_path):
        print("❌ File không tồn tại!")
        return

    # Menu
    print("\n🎯 CHỌN CHỨC NĂNG:")
    print("1. Xem trước (không lưu)")
    print("2. Trích xuất ngay")
    print("3. Xem trước rồi quyết định")
    
    choice = input("Chọn (1-3): ").strip()
    
    if choice == "1":
        preview_images_and_indices(docx_path)
        
    elif choice == "2":
        extract_images_with_precise_index(docx_path)
        
    elif choice == "3":
        preview_images_and_indices(docx_path)
        confirm = input("\n❓ Tiếp tục trích xuất? (y/n): ").strip().lower()
        if confirm == 'y':
            extract_images_with_precise_index(docx_path)
        else:
            print("✋ Hủy trích xuất.")
    else:
        print("❌ Lựa chọn không hợp lệ!")

if __name__ == "__main__":
    main()